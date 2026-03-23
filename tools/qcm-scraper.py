"""
QCM Scraper — prepa.concours-sesame.net
Scrape les 6 matières en parallèle, exporte en .docx
"""
import os, sys, json, time, re, getpass
from datetime import date
from concurrent.futures import ThreadPoolExecutor, as_completed

# ─── CHEMINS ────────────────────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR    = os.path.join(BASE_DIR, "sesame-prep", "annales", "QCM")
os.makedirs(OUT_DIR, exist_ok=True)

LOGIN_URL  = "https://connexion.concours-sesame.net/"
RETURN_URL = "https://prepa.concours-sesame.net"

SUBJECTS = [
    ("Enjeux contemporains",  "https://prepa.concours-sesame.net/QCM-enjeux-contemporains"),
    ("Mathématiques",         "https://prepa.concours-sesame.net/QCM-Mathematiques"),
    ("Compétences digitales", "https://prepa.concours-sesame.net/QCM-competences-digitales"),
    ("Français",              "https://prepa.concours-sesame.net/QCM-francais"),
    ("Langues étrangères",    "https://prepa.concours-sesame.net/QCM-langues-etrangeres"),
    ("Analyse documentaire",  "https://prepa.concours-sesame.net/QCM-analyse-documentaire"),
]

# ─── CREDENTIALS ────────────────────────────────────────────────────────────
def get_credentials():
    env_file = os.path.join(BASE_DIR, ".env")
    email, pwd = None, None
    if os.path.exists(env_file):
        for line in open(env_file, encoding="utf-8"):
            if line.startswith("SESAME_EMAIL="):
                email = line.split("=", 1)[1].strip()
            if line.startswith("SESAME_PASSWORD="):
                pwd = line.split("=", 1)[1].strip()
    if not email:
        email = input("Email prepa.concours-sesame.net : ").strip()
    if not pwd:
        pwd = getpass.getpass("Mot de passe : ")
    return email, pwd

# ─── LOGIN ──────────────────────────────────────────────────────────────────
def login_requests(email, pwd):
    """Connexion via requests.Session — retourne la session ou None si échec."""
    import requests
    s = requests.Session()
    s.headers.update({"User-Agent": "Mozilla/5.0"})
    # Récupérer le token CSRF si présent
    r = s.get("https://prepa.concours-sesame.net/", timeout=10)
    csrf = re.search(r'name="__RequestVerificationToken"[^>]+value="([^"]+)"', r.text)
    payload = {
        "Username": email,
        "Password": pwd,
        "returnUrl": RETURN_URL,
        "errorUrl":  RETURN_URL + "/Identity/Account/Login",
    }
    if csrf:
        payload["__RequestVerificationToken"] = csrf.group(1)
    r2 = s.post(LOGIN_URL, data=payload, timeout=10, allow_redirects=True)
    # Vérifier qu'on est connecté (pas redirigé vers login)
    if "connexion" in r2.url.lower() or "login" in r2.url.lower():
        return None
    return s

# ─── SCRAPE STATIQUE ────────────────────────────────────────────────────────
def scrape_static(session, name, url):
    """Tente de parser les questions depuis le HTML brut."""
    from bs4 import BeautifulSoup
    r = session.get(url, timeout=15)
    soup = BeautifulSoup(r.text, "html.parser")

    questions = []
    # Patterns courants pour les quiz HTML
    q_blocks = (soup.find_all(class_=re.compile(r"question", re.I)) or
                soup.find_all("li", class_=re.compile(r"q\d|quiz", re.I)))

    for block in q_blocks:
        text = block.get_text(separator=" ", strip=True)
        if len(text) < 20:
            continue
        questions.append({"text": text, "options": [], "correct": "?", "explanation": ""})

    return questions if questions else None  # None = fallback Selenium

# ─── SCRAPE SELENIUM ────────────────────────────────────────────────────────
def scrape_selenium(email, pwd, name, url):
    """Selenium headless — réutilise les cookies de login."""
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC

    opts = Options()
    opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--window-size=1280,900")

    driver = webdriver.Chrome(options=opts)
    wait   = WebDriverWait(driver, 12)
    questions = []

    try:
        # Login
        driver.get("https://prepa.concours-sesame.net/")
        wait.until(EC.presence_of_element_located((By.NAME, "Username")))
        driver.find_element(By.NAME, "Username").send_keys(email)
        driver.find_element(By.NAME, "Password").send_keys(pwd)
        driver.find_element(By.CSS_SELECTOR, "button[type=submit], input[type=submit]").click()
        time.sleep(2)

        # Aller sur le QCM
        driver.get(url)
        time.sleep(2)

        # Cliquer sur "Commencer" si présent
        for sel in ["button.start", ".btn-start", "a.commencer", "[class*='start']", "button"]:
            try:
                btns = driver.find_elements(By.CSS_SELECTOR, sel)
                for btn in btns:
                    if any(k in btn.text.lower() for k in ["commencer", "démarrer", "start", "▶", "jouer"]):
                        btn.click()
                        time.sleep(1.5)
                        break
            except:
                pass

        # Boucle sur les questions
        for q_num in range(1, 11):
            try:
                # Attendre la question
                time.sleep(1)
                page_src = driver.page_source
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(page_src, "html.parser")

                # Extraire texte question
                q_text = ""
                for sel in [".question-text", ".question", "h2", "h3", ".quiz-question"]:
                    el = soup.select_one(sel)
                    if el and len(el.get_text(strip=True)) > 15:
                        q_text = el.get_text(strip=True)
                        break

                # Extraire options
                options = []
                for sel in [".answer", ".option", ".choice", "li.reponse", ".reponse", "label"]:
                    els = soup.select(sel)
                    if len(els) >= 2:
                        options = [e.get_text(strip=True) for e in els if e.get_text(strip=True)]
                        break

                # Cliquer première option disponible
                clicked = False
                for sel in [".answer", ".option", ".choice", "li.reponse", ".reponse", "label", "input[type=radio]"]:
                    try:
                        clickable = driver.find_elements(By.CSS_SELECTOR, sel)
                        if clickable:
                            driver.execute_script("arguments[0].click();", clickable[0])
                            clicked = True
                            time.sleep(1.5)
                            break
                    except:
                        pass

                # Lire la correction
                page_src2 = driver.page_source
                soup2 = BeautifulSoup(page_src2, "html.parser")

                correct = "?"
                explanation = ""
                for sel in [".correct-answer", ".bonne-reponse", ".answer.correct", "[class*='correct']"]:
                    el = soup2.select_one(sel)
                    if el:
                        correct = el.get_text(strip=True)
                        break
                for sel in [".explanation", ".explication", ".feedback", "[class*='explication']", "[class*='explanation']"]:
                    el = soup2.select_one(sel)
                    if el:
                        explanation = el.get_text(strip=True)
                        break

                questions.append({
                    "num": q_num,
                    "text": q_text or f"Question {q_num} (non capturée)",
                    "options": options[:4],
                    "correct": correct,
                    "explanation": explanation,
                })
                print(f"  [{name}] Q{q_num}/10 ✓")

                # Question suivante
                for sel in ["button.next", ".btn-next", "[class*='next']", "[class*='suivant']", "button", "a"]:
                    try:
                        btns = driver.find_elements(By.CSS_SELECTOR, sel)
                        for btn in btns:
                            if any(k in btn.text.lower() for k in ["suivant", "next", "continuer", "suite"]):
                                driver.execute_script("arguments[0].click();", btn)
                                break
                    except:
                        pass

            except Exception as e:
                print(f"  [{name}] Q{q_num} erreur: {e}")
                questions.append({"num": q_num, "text": f"[données partielles — Q{q_num}]",
                                  "options": [], "correct": "?", "explanation": ""})

    finally:
        driver.quit()

    return questions

# ─── WORKER ─────────────────────────────────────────────────────────────────
def scrape_subject(args):
    email, pwd, name, url = args
    print(f"⟳ Démarrage : {name}")
    # Essai statique d'abord
    try:
        import requests
        s = login_requests(email, pwd)
        if s:
            qs = scrape_static(s, name, url)
            if qs:
                print(f"✓ {name} — {len(qs)} questions (statique)")
                return name, qs
    except Exception as e:
        print(f"  [{name}] statique échoué ({e}), bascule Selenium")

    # Fallback Selenium
    try:
        qs = scrape_selenium(email, pwd, name, url)
        print(f"✓ {name} — {len(qs)} questions (Selenium)")
        return name, qs
    except Exception as e:
        print(f"✗ {name} — ÉCHEC : {e}")
        return name, []

# ─── EXPORT WORD ────────────────────────────────────────────────────────────
def export_docx(all_data):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    today = date.today().strftime("%Y-%m-%d")

    # Titre
    t = doc.add_heading(f"QCM Sésame — {today}", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, questions in all_data:
        doc.add_heading(name, level=1)
        if not questions:
            doc.add_paragraph("⚠ Aucune question capturée pour cette matière.")
            continue

        for q in questions:
            num = q.get("num", "?")
            # Question
            p = doc.add_paragraph()
            run = p.add_run(f"Q{num}. {q['text']}")
            run.bold = True

            # Options
            for i, opt in enumerate(q.get("options", [])):
                letters = "ABCD"
                doc.add_paragraph(f"   {letters[i] if i < 4 else '?'}. {opt}")

            # Bonne réponse
            rep_p = doc.add_paragraph()
            rep_run = rep_p.add_run(f"✅ Bonne réponse : {q['correct']}")
            rep_run.bold = True
            rep_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

            # Explication
            if q.get("explanation"):
                exp_p = doc.add_paragraph(f"📝 {q['explanation']}")
                exp_p.paragraph_format.left_indent = Pt(20)

            doc.add_paragraph("─" * 60)

    out_path = os.path.join(OUT_DIR, f"QCM_Sesame_{today}.docx")
    doc.save(out_path)
    return out_path

# ─── MAIN ───────────────────────────────────────────────────────────────────
def main():
    print("=" * 50)
    print("  QCM Scraper — Sésame (parallèle)")
    print("=" * 50)

    email, pwd = get_credentials()

    # Vérifier login
    print("\n🔐 Connexion...")
    import requests
    s = login_requests(email, pwd)
    if not s:
        print("⚠ Login requests échoué — Selenium gérera l'auth dans chaque thread.")

    # Scrape en parallèle (6 workers)
    args = [(email, pwd, name, url) for name, url in SUBJECTS]
    results_dict = {}

    print(f"\n🚀 Scraping des {len(SUBJECTS)} matières en parallèle...\n")
    t0 = time.time()

    with ThreadPoolExecutor(max_workers=6) as executor:
        futures = {executor.submit(scrape_subject, a): a[2] for a in args}
        for future in as_completed(futures):
            name, questions = future.result()
            results_dict[name] = questions

    elapsed = time.time() - t0
    print(f"\n⏱ Scraping terminé en {elapsed:.1f}s")

    # Reconstruire dans l'ordre original
    all_data = [(name, results_dict.get(name, [])) for name, _ in SUBJECTS]
    total_q = sum(len(qs) for _, qs in all_data)
    print(f"📊 {total_q} questions capturées au total")

    # Export
    print("\n📄 Export Word...")
    path = export_docx(all_data)
    print(f"✅ Fichier créé : {path}")

    # Sauvegarder aussi en JSON pour question-bank
    json_path = os.path.join(OUT_DIR, f"QCM_Sesame_{date.today().strftime('%Y-%m-%d')}.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({name: qs for name, qs in all_data}, f, ensure_ascii=False, indent=2)
    print(f"✅ JSON sauvegardé : {json_path}")

if __name__ == "__main__":
    main()
